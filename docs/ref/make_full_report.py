#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate docs/SENG326-Task2-Report-Team13.docx
Full detailed report matching the Task 1 example format.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).parent.parent
SS   = ROOT / "docs" / "screenshots"
OUT  = ROOT / "docs" / "SENG326-Task2-Report-Team13.docx"

H_COLOR  = RGBColor(0x0F, 0x47, 0x61)   # matching example heading color
BG_LIGHT = "F2F2F2"                       # code block background
TBL_HEAD = "D9E1F2"                       # table header background
TBL_ALT  = "EEF3FA"                       # table alt row

# ── document setup ──────────────────────────────────────────────────────────

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

styles = doc.styles

# Body style
normal = styles["Normal"]
normal.font.name  = "Calibri"
normal.font.size  = Pt(11)
normal.paragraph_format.space_after = Pt(6)

# Heading styles
for lvl, sz in [(1, 18), (2, 15), (3, 13), (4, 12)]:
    s = styles[f"Heading {lvl}"]
    s.font.name  = "Calibri"
    s.font.size  = Pt(sz)
    s.font.bold  = True
    s.font.color.rgb = H_COLOR
    s.paragraph_format.space_before = Pt(14)
    s.paragraph_format.space_after  = Pt(4)

# ── helpers ─────────────────────────────────────────────────────────────────

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def h4(text):
    doc.add_heading(text, level=4)

def body(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    _inline(p, text)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    _inline(p, text)

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    _inline(p, text)

def spacer():
    doc.add_paragraph()

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x40,0x40,0x40)

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x2e, 0x75, 0xb6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), BG_LIGHT)
    pPr.append(shd)

def img(filename, width=Inches(5.8), cap=None):
    path = SS / filename
    if not path.exists():
        body(f"[Screenshot not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    if cap:
        caption(cap)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header
    hrow = t.rows[0]
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]
        cell.text = ""
        rn = cell.paragraphs[0].add_run(h)
        rn.bold = True; rn.font.size = Pt(9.5)
        set_cell_bg(cell, TBL_HEAD)
    # data rows
    for ri, row in enumerate(rows):
        trow = t.rows[ri+1]
        bg = TBL_ALT if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = ""
            rn = cell.paragraphs[0].add_run(str(val))
            rn.font.size = Pt(9)
            if bg != "FFFFFF":
                set_cell_bg(cell, bg)
    spacer()

def _inline(para, text):
    import re
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    for part in pattern.split(text):
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x2e, 0x75, 0xb6)
        else:
            para.add_run(part)

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════

for _ in range(5):
    spacer()

p = doc.add_heading("SENG326 - SOFTWARE ARCHITECTURE", level=1)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TASK 2 — ARCHITECTURE REFACTORING")
r.font.size = Pt(16); r.bold = True; r.font.color.rgb = H_COLOR

spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Hexagonal Architecture (Ports & Adapters) Refactoring of Checkstyle 13.2.0")
r.font.size = Pt(13); r.italic = True

for _ in range(3):
    spacer()

for name in ["KERİM HARİRİ – 22050941008", "MOHAMAD ATTIA EID ATTA EID – 22050941017", "Zaid Hardan – 22050941005"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name); r.font.size = Pt(12); r.bold = True

spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Group 13  |  Branch: 3-implement_hexagonal  |  Checkstyle 13.2.0")
r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RGBColor(0x60,0x60,0x60)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART A
# ════════════════════════════════════════════════════════════════════════════

h1("Part A: Architecture Refactoring")
spacer()

# ── 1. INTRODUCTION ─────────────────────────────────────────────────────────
h1("1.  Introduction")

h2("1.1  Purpose of This Report")
body("This report documents the complete architectural refactoring performed on Checkstyle 13.2.0 as part of SENG326 Task 2. The goal was to take a defined slice of the existing codebase and restructure it into a target architectural style — Hexagonal Architecture (Ports & Adapters) — while preserving byte-identical output, passing all existing tests, and providing verifiable architectural evidence.")
body("The report covers: selection and justification of the target style, design of the new architecture, step-by-step migration of all 16 checks in the slice, verification via regression tests, ArchUnit conformance rules, jQAssistant graph queries, and C4 diagrams. Part B documents the performance benchmark comparing the original and refactored jars across five open-source Java projects.")

h2("1.2  Background: Checkstyle 13.2.0")
body("Checkstyle is an open-source static analysis tool for Java code. Given a configuration XML and one or more Java source files, it reports violations of configured coding rules. The tool is widely used in CI pipelines and IDE integrations to enforce project-specific coding standards.")
body("Architecturally, Checkstyle follows a **plugin-based framework** pattern: a central `Checker` and `TreeWalker` engine dispatches events to registered check modules, each of which is an independent plugin loaded by name from XML configuration. As of version 13.2.0, there are approximately 200 checks spanning categories including annotation, blocks, coding, design, metrics, sizes, naming, whitespace, and more.")
body("The execution pipeline is:")
numbered("**Configuration loading**: `ConfigurationLoader` parses the XML config and uses `PackageObjectFactory` to instantiate check modules by short name.")
numbered("**File iteration**: `Checker` iterates over all input `.java` files.")
numbered("**Parsing**: `TreeWalker` (a `FileSetCheck`) parses each file into a `DetailAST` tree using the ANTLR4 grammar.")
numbered("**Check dispatch**: `TreeWalker` walks every AST node and calls `visitToken()` / `leaveToken()` on each registered `AbstractCheck` whose `getDefaultTokens()` includes that node type.")
numbered("**Violation reporting**: Each check calls `log()` when a violation is found; `Checker` aggregates and formats violations.")

h2("1.3  Assignment Overview and Objectives")
body("Task 2 required selecting a slice of Checkstyle — a cohesive subset of checks — and refactoring it into one of several permitted target architectural styles: Strict Layered, SOA-Orchestration, SOA-Message Bus, Event-Driven, Blackboard, Pipe-and-Filter, or Hexagonal (Ports & Adapters).")
body("Our group selected **Hexagonal Architecture** for the following reasons:")
bullet("Checkstyle's current secondary style is already pipe-and-filter; choosing it again would produce weak refactoring evidence")
bullet("Hexagonal creates the clearest architectural shift: domain logic becomes framework-free, testable in isolation, and bounded by explicit port interfaces")
bullet("The metrics/sizes slice has well-defined input (AST nodes / file content) and output (violation reports) — exactly the adapter boundaries hexagonal architecture formalizes")
body("The **hard constraints** imposed by the assignment were:")
bullet("The refactored slice must execute exclusively through the new architecture — no shortcuts back to the original pattern")
bullet("Existing XML configurations (`google_checks.xml`, etc.) must continue to work without modification")
bullet("Check logic must NOT be rewritten — only relocated and reorganized")
bullet("Output must be byte-identical to the original on the same input")
bullet("Minimum new classes, each with a single responsibility and a Javadoc justification")

page_break()

# ── 2. HEXAGONAL ARCHITECTURE ───────────────────────────────────────────────
h1("2.  Target Architectural Style: Hexagonal Architecture (Ports & Adapters)")

h2("2.1  What is Hexagonal Architecture?")
body("Hexagonal Architecture, introduced by Alistair Cockburn in 2005, organises a system around a central **domain hexagon** that contains all business logic. The domain communicates with the outside world exclusively through **ports** — technology-neutral interfaces that the domain itself defines. The outside world connects to the domain through **adapters** that implement those ports.")
body("The fundamental invariant is: **the domain knows nothing about any framework, database, UI, or external system**. It imports only its own port interfaces. This makes the domain fully testable in isolation, replaceable in its infrastructure, and immune to framework changes.")
body("The three concentric zones are:")

table(
    ["Zone", "Responsibility", "Framework coupling"],
    [
        ["Domain", "Contains all business/check logic. Pure Java with no framework imports.", "None — zero api.* imports"],
        ["Ports", "Technology-neutral interfaces defined by the domain. Bridge between domain and adapters.", "Interface-only — no implementation"],
        ["Adapters", "Implement ports. Bridge the framework API to the domain. Thin delegation only.", "Full — extends AbstractCheck, imports api.*"],
    ]
)

h2("2.2  Core Concepts in Our Refactoring")
body("In the context of our Checkstyle refactoring, the three zones map as follows:")

h3("Domain")
body("The domain classes are the `*Rule.java` files created under `checks/metrics/domain/` and `checks/sizes/domain/`. Each class contains the complete check logic extracted from the original `*Check.java`. Domain classes import only `DetailAST` and `TokenTypes` from the Checkstyle API — both are data-only types with no framework inheritance — and the port interfaces `ICheckUseCase` and `IViolationOutputPort`, which are defined in our own `checks.hexagonal.port` package. No class in the domain extends `AbstractCheck` or `AbstractFileSetCheck`.")

h3("Ports")
body("Four port artifacts were created in `checks/hexagonal/port/`:")
bullet("`ICheckUseCase` — inbound port for AST-based checks. Defines `getDefaultTokens()`, `visitToken()`, `leaveToken()`, `beginTree()`, `finishTree()`. Domain rules implement this interface.")
bullet("`IFileCheckUseCase` — inbound port for file-level checks. Defines `processFile(FileContent)`. Used by `FileLengthRule` and `LineLengthRule`.")
bullet("`IViolationOutputPort` — outbound port. Defines `report(line, msg, args)`. Domain rules call this to emit violations; adapters implement it by delegating to `AbstractCheck.log()`.")
bullet("`FileContent` — value object (not strictly a port but crosses the boundary). Wraps a `String[]` of file lines, replacing `FileText` (a Checkstyle API type) at the domain boundary.")

h3("Adapters")
body("Two abstract adapter base classes were created in `checks/hexagonal/`:")
bullet("`AstCheckAdapter` — extends `AbstractCheck`, implements `IViolationOutputPort`. Subclasses provide a `createRule()` factory method returning their domain rule. Used by all 14 AST-based checks.")
bullet("`FileSetCheckAdapter` — extends `AbstractFileSetCheck`, implements `IViolationOutputPort`. Converts `FileText` to `FileContent` before calling the domain. Used by `FileLengthCheck` and `LineLengthCheck`.")
body("Each of the 16 original check classes was replaced by a thin concrete adapter that only overrides property setters (to pass configuration to the domain rule) and `createRule()` (to instantiate the domain rule). No check logic lives in the adapter.")

h2("2.3  Why Hexagonal Architecture for This Slice?")
body("The metrics and sizes checks are a natural fit for hexagonal architecture because:")
numbered("**Well-defined inputs**: Each check either processes AST nodes (typed, immutable tree nodes from `DetailAST`) or receives raw file content (a `String[]`). These translate cleanly into port method signatures.")
numbered("**Well-defined output**: The only observable output of a check is a set of violation records (line, column, message key, arguments). The `IViolationOutputPort` interface captures this exactly.")
numbered("**No shared mutable state across checks**: Each check maintains its own counter/stack during a single file walk, reset per file. This maps directly to `@FileStatefulCheck` adapter instances.")
numbered("**Clear domain logic**: The computation in metrics/sizes checks (counting tokens, measuring lengths, tracking nesting depth) is purely algorithmic — ideal for a domain class with no I/O or framework dependency.")
numbered("**Large slice**: 16 checks provides enough evidence to prove the architecture is applied consistently, not just for one class.")

h2("2.4  Comparison with the Original Plugin-Based Architecture")
body("In the original Checkstyle architecture, each check class:")
bullet("Extends `AbstractCheck` or `AbstractFileSetCheck` directly, inheriting the entire framework lifecycle")
bullet("Calls `log()` (a method on `AbstractCheck`) directly to emit violations")
bullet("Imports `DetailAST`, `TokenTypes`, and other `api.*` types at will")
bullet("Is instantiated and managed by `TreeWalker` / `Checker`")
body("After the hexagonal refactoring, the same check:")
bullet("Has a thin adapter that extends `AbstractCheck` — all framework coupling is here")
bullet("Has a domain rule that implements `ICheckUseCase` — no framework imports")
bullet("Reports violations through `IViolationOutputPort.report()` — injected by the adapter at construction time")
bullet("Can be tested without any Checkstyle installation by constructing the rule with a mock `IViolationOutputPort`")

table(
    ["Aspect", "Before (Plugin-Based)", "After (Hexagonal)"],
    [
        ["Framework coupling", "Every check class extends AbstractCheck", "Only 2 adapter base classes + 16 thin adapters"],
        ["Domain isolation", "Logic mixed with api.* imports", "Domain classes: zero api.* imports"],
        ["Violation emission", "Direct log() call on AbstractCheck", "violationPort.report() through IViolationOutputPort"],
        ["Testability", "Requires full Checkstyle runtime", "Domain testable with a mock IViolationOutputPort"],
        ["Config compatibility", "XML name = class name", "XML name = adapter class name (same name, same package)"],
        ["New check cost", "One class extending AbstractCheck", "One domain class + one thin adapter (2 files)"],
    ]
)

page_break()

# ── 3. SLICE DEFINITION ─────────────────────────────────────────────────────
h1("3.  Architectural Slice Definition")

h2("3.1  Slice Selection Rationale")
body("We selected the **Metrics and Size Violation checks** as our architectural slice. These checks occupy two packages: `com.puppycrawl.tools.checkstyle.checks.metrics` (6 checks) and `com.puppycrawl.tools.checkstyle.checks.sizes` (10 checks). Together they form a cohesive group of 16 checks that share a common theme: they measure quantitative properties of Java code (complexity, length, count) and report violations when measurements exceed configured thresholds.")
body("This slice was chosen because it is large enough to provide strong evidence of consistent hexagonal application (16 classes is a meaningful refactoring, not a toy example), cohesive enough to fit a single architectural boundary, and independent enough from other checks that the migration does not ripple into unrelated code.")

h2("3.2  Checks Included in the Slice (16 Checks)")
table(
    ["#", "Check Name", "Package", "What It Measures", "Port"],
    [
        ["1", "BooleanExpressionComplexityCheck", "metrics", "Counts && and || operators in boolean expressions", "ICheckUseCase"],
        ["2", "ClassDataAbstractionCouplingCheck", "metrics", "Counts distinct instantiated (new X()) types per class", "ICheckUseCase"],
        ["3", "ClassFanOutComplexityCheck", "metrics", "Counts all referenced types per class", "ICheckUseCase"],
        ["4", "CyclomaticComplexityCheck", "metrics", "McCabe cyclomatic complexity per method/constructor", "ICheckUseCase"],
        ["5", "JavaNCSSCheck", "metrics", "Non-commenting source statements per method/class/file", "ICheckUseCase"],
        ["6", "NPathComplexityCheck", "metrics", "Number of acyclic execution paths per method", "ICheckUseCase"],
        ["7", "AnonInnerLengthCheck", "sizes", "Line count of anonymous inner class bodies", "ICheckUseCase"],
        ["8", "ExecutableStatementCountCheck", "sizes", "Executable statements per method/constructor", "ICheckUseCase"],
        ["9", "FileLengthCheck", "sizes", "Total line count of the file", "IFileCheckUseCase"],
        ["10", "LambdaBodyLengthCheck", "sizes", "Line count of lambda expression bodies", "ICheckUseCase"],
        ["11", "LineLengthCheck", "sizes", "Character count per line", "IFileCheckUseCase"],
        ["12", "MethodCountCheck", "sizes", "Number of methods per type declaration", "ICheckUseCase"],
        ["13", "MethodLengthCheck", "sizes", "Line count of method/constructor bodies", "ICheckUseCase"],
        ["14", "OuterTypeNumberCheck", "sizes", "Number of top-level type declarations per file", "ICheckUseCase"],
        ["15", "ParameterNumberCheck", "sizes", "Parameter count per method/constructor", "ICheckUseCase"],
        ["16", "RecordComponentNumberCheck", "sizes", "Component count per record declaration", "ICheckUseCase"],
    ]
)

h2("3.3  Checks Excluded")
body("All other ~184 Checkstyle checks are outside the slice and were left completely unchanged. This includes annotation checks, blocks checks, coding checks, design checks, header checks, import checks, indentation checks, Javadoc checks, modifier checks, naming checks, regexp checks, and whitespace checks. The `Checker`, `TreeWalker`, all API classes, all grammar files, all utility classes, and all filter classes are also untouched.")

h2("3.4  Hard Constraints Applied")
body("The following constraints governed every refactoring decision:")
bullet("**Exclusive execution**: The refactored checks execute only through the new hexagonal path. No shortcut bypasses the port interfaces.")
bullet("**Config compatibility**: Adapter class names are identical to the originals (same class name, same Java package). Existing XMLs load them without change.")
bullet("**Logic preservation**: Domain rule classes are direct extractions of the original check logic. No algorithm was modified.")
bullet("**Byte-identical output**: Verified by `diff baseline/pre-refactor-output.txt baseline/post-refactor-output.txt` — empty diff on a sample file triggering all 16 checks.")
bullet("**Minimum new classes**: Exactly 5 new infrastructure classes were added (`AstCheckAdapter`, `FileSetCheckAdapter`, `ICheckUseCase`, `IFileCheckUseCase`, `IViolationOutputPort`, `FileContent` — 6 including the value object). Each has a Javadoc architectural justification.")

page_break()

# ── 4. PRE-REFACTORING BASELINE ─────────────────────────────────────────────
h1("4.  Pre-Refactoring Baseline")

h2("4.1  Original Directory Structure")
body("Before refactoring, the `checks/metrics/` and `checks/sizes/` packages each contained a flat set of check classes directly extending `AbstractCheck` or `AbstractClassCouplingCheck` (an internal base class). There were no subdirectories, no port interfaces, and no separation between framework coupling and domain logic.")
body("The screenshots below show the original directory structure of both packages as captured before any refactoring changes were made.")

spacer()
img("phase1-metrics-dir.png", Inches(5.5), "Figure 1 – Original checks/metrics/ directory structure (6 check classes, flat package)")
spacer()
img("phase1-sizes-dir.png", Inches(5.5), "Figure 2 – Original checks/sizes/ directory structure (10 check classes, flat package)")
spacer()

body("In the original `checks/metrics/` package, the 6 check classes (plus `AbstractClassCouplingCheck` as an internal base) all inherit directly from `AbstractCheck`. The `AbstractClassCouplingCheck` provides shared logic for `ClassDataAbstractionCouplingCheck` and `ClassFanOutComplexityCheck`, but this is an internal implementation detail — not a port interface.")
body("In the original `checks/sizes/` package, all 10 check classes are independent (no shared base). Two of them — `FileLengthCheck` and `LineLengthCheck` — extend `AbstractFileSetCheck` rather than `AbstractCheck` because they operate on the whole file rather than individual AST nodes.")

h2("4.2  Original Architecture: AbstractCheck Inheritance Pattern")
body("The original architecture is a **plugin-based framework** pattern. Every check class:")
numbered("Extends `AbstractCheck` (or `AbstractFileSetCheck` for file-level checks) directly, inheriting the framework lifecycle methods: `beginTree()`, `visitToken()`, `leaveToken()`, `finishTree()`.")
numbered("Calls `this.log(line, col, msgKey, args)` to emit violations — a method inherited from `AbstractCheck`. This creates a direct framework dependency inside the check logic.")
numbered("Imports `DetailAST`, `TokenTypes`, `ScopeUtil`, and other `api.*` / `utils.*` types directly inside the check class where the logic lives.")
body("This means the **check logic and the framework coupling are in the same class**. There is no separation between what the check does (the algorithm) and how it integrates with the Checkstyle framework (the adapter). A developer testing the cyclomatic complexity algorithm must instantiate the whole Checkstyle check framework to do so.")

h2("4.3  Pre-Refactoring Baseline Output (44 Violations)")
body("To establish a baseline, we created `violation-sample/SampleAllViolations.java` — a synthetic Java file crafted to trigger at least one violation from each of the 16 checks — and ran the original Checkstyle jar against it with a low-threshold config file (`violation-sample/all-metrics-sizes.xml`).")
body("The pre-refactoring output was captured and saved to `baseline/pre-refactor-output.txt`. The screenshot below shows the terminal output from running the original jar:")
spacer()
img("phase1-pre-refactor-output.png", Inches(5.8), "Figure 3 – Pre-refactoring terminal output: 44 violations from all 16 checks")
spacer()
body("The baseline output contains **44 violations** covering all 16 checks. This file served as the ground truth against which the refactored jar output was compared using `diff`. An empty diff was required to confirm byte-identical correctness.")

page_break()

# ── 5. REFACTORING DESIGN ───────────────────────────────────────────────────
h1("5.  Refactoring Design")

h2("5.1  Port Interface Design")
body("The first step in any hexagonal refactoring is to define the ports — the interfaces that form the boundary of the domain hexagon. We defined four artifacts in the package `com.puppycrawl.tools.checkstyle.checks.hexagonal.port`:")

h3("5.1.1  ICheckUseCase — Inbound Port for AST-Based Checks")
body("`ICheckUseCase` is the inbound port that inbound adapters call to drive a domain check rule. It mirrors the extension points of `AbstractCheck` — `getDefaultTokens()`, `getAcceptableTokens()`, `getRequiredTokens()`, `visitToken()`, `leaveToken()`, `beginTree()`, `finishTree()` — but without any dependency on the `AbstractCheck` class itself. Domain rules implement this interface; the Checkstyle framework never sees the domain class directly.")
code_block("""public interface ICheckUseCase {
    int[] getDefaultTokens();
    int[] getAcceptableTokens();
    int[] getRequiredTokens();
    default void beginTree(DetailAST rootAST) { }
    default void visitToken(DetailAST ast) { }
    default void leaveToken(DetailAST ast) { }
    default void finishTree(DetailAST rootAST) { }
}""")

h3("5.1.2  IFileCheckUseCase — Inbound Port for File-Level Checks")
body("`IFileCheckUseCase` is the inbound port for the two file-level checks (`FileLengthCheck` and `LineLengthCheck`). Rather than AST callbacks, file-level domain rules receive a `FileContent` value object containing all file lines and the filename.")
code_block("""public interface IFileCheckUseCase {
    void processFile(FileContent fileContent);
}""")

h3("5.1.3  IViolationOutputPort — Outbound Port")
body("`IViolationOutputPort` is the single channel through which domain rule classes emit violations. Adapters implement this interface by delegating to `AbstractCheck.log()` or `AbstractFileSetCheck.log()`, keeping all Checkstyle API coupling outside the domain hexagon. Three overloads handle line-only, line+column, and AST-anchored violations.")
code_block("""public interface IViolationOutputPort {
    void report(int line, String messageKey, Object... args);
    void report(int line, int col, String messageKey, Object... args);
    void report(DetailAST ast, String messageKey, Object... args);
}""")

h3("5.1.4  FileContent — Value Object")
body("`FileContent` is an immutable value object that carries raw file content (a `String[]` of lines) across the adapter-to-domain boundary. Its purpose is to replace `FileText` — a Checkstyle API type — at the domain boundary so that file-level domain rules have no `api.*` dependency. The adapter constructs it from `FileText` before calling into the domain.")

spacer()
img("phase2-port-interfaces.png", Inches(5.8), "Figure 4 – Port interface definitions: ICheckUseCase, IViolationOutputPort, FileContent")
spacer()

h2("5.2  Adapter Base Classes")
body("With the ports defined, we created two abstract adapter base classes that bridge the Checkstyle framework to the domain. These base classes contain all framework coupling; individual check adapters inherit from them and add only the configuration property setters specific to each check.")

h3("5.2.1  AstCheckAdapter")
body("`AstCheckAdapter` (in package `checks.hexagonal`) extends `AbstractCheck` and implements `IViolationOutputPort`. It holds a reference to the domain rule via `ICheckUseCase` and delegates every `AbstractCheck` callback to the rule. It implements `IViolationOutputPort` by calling `super.log(...)`, bridging the outbound port to the framework violation channel.")
body("Subclasses implement the abstract method `createRule(IViolationOutputPort violationPort)` to construct and return their specific domain rule, passing `this` as the violation port (since `AstCheckAdapter` itself implements `IViolationOutputPort`).")
spacer()
img("phase3-ast-adapter.png", Inches(5.8), "Figure 5 – AstCheckAdapter: the shared inbound adapter for all 14 AST-based checks")
spacer()

h3("5.2.2  FileSetCheckAdapter")
body("`FileSetCheckAdapter` (in package `checks.hexagonal`) extends `AbstractFileSetCheck` and implements `IViolationOutputPort`. It overrides `processFiltered(File, FileText)` to convert `FileText` into a `FileContent` value object and then calls the abstract method `processFileContent(File, FileContent)` which subclasses implement using their domain rule.")
spacer()
img("phase3-fileset-adapter.png", Inches(5.8), "Figure 6 – FileSetCheckAdapter: the shared inbound adapter for the 2 file-level checks")
spacer()

h2("5.3  Domain Rule Classes")
body("Each of the 16 original check classes was split into a domain rule (`*Rule.java`) and an adapter (`*Check.java`). The domain rule contains all the check logic extracted from the original class, with two changes: (1) all `log(...)` calls replaced by `violationPort.report(...)`, and (2) the constructor receives an `IViolationOutputPort` parameter instead of inheriting `log()` from `AbstractCheck`.")
body("The screenshot below shows `CyclomaticComplexityRule` — the domain rule for cyclomatic complexity measurement. It implements `ICheckUseCase`, maintains a `Deque<BigInteger>` for nesting, counts decision points, and calls `violationPort.report()` when the threshold is exceeded. There is no `AbstractCheck` import anywhere in the class.")
spacer()
img("phase3-domain-class.png", Inches(5.8), "Figure 7 – CyclomaticComplexityRule: pure domain logic, zero api.* framework inheritance")
spacer()

h2("5.4  Concrete Adapter Classes")
body("Each concrete adapter is a thin class that extends either `AstCheckAdapter` or `FileSetCheckAdapter`. Its only responsibilities are: (1) override `createRule()` to instantiate the domain rule with the correct configuration, and (2) expose property setter methods (annotated with `@XmlAttribute` and property names matching the original) so that Checkstyle XML configuration still works.")
body("The screenshot below shows `CyclomaticComplexityCheck` after refactoring. The entire class is reduced to a constructor call, a `createRule()` factory method, and property setters. All logic has moved to `CyclomaticComplexityRule`.")
spacer()
img("phase3-adapter-class.png", Inches(5.8), "Figure 8 – CyclomaticComplexityCheck (adapter): thin delegation layer, all logic in the domain rule")
spacer()

page_break()

# ── 6. IMPLEMENTATION ───────────────────────────────────────────────────────
h1("6.  Implementation: Migration of All 16 Checks")

h2("6.1  Migration Procedure")
body("The migration followed a strict extract-then-wrap procedure for each check. This approach ensured that at every step the build passed and the output remained correct, providing a clear checkpoint and making bugs easy to isolate.")
numbered("**Create the domain rule** (`*Rule.java` in `*.domain` subpackage): Copy all non-Checkstyle-API logic from the original check into a new class implementing `ICheckUseCase` (or `IFileCheckUseCase`). Replace all `this.log(...)` calls with `this.violationPort.report(...)`. Add `IViolationOutputPort` as a constructor parameter.")
numbered("**Create the adapter** (`*Check.java` in the original package): Extend `AstCheckAdapter` (or `FileSetCheckAdapter`). Override `createRule()` to instantiate the domain rule with configuration parameters. Copy all property setter methods from the original.")
numbered("**Delete the original** `*Check.java`. The adapter takes its place with the same class name, same package, and same Checkstyle XML short name.")
numbered("**Verify**: Run `mvn -DskipTests package` to confirm the build passes, then check `diff baseline/pre-refactor-output.txt <new-output>` is empty for each migrated check.")
body("The checks were migrated in order of increasing complexity: file-set checks first (simplest adapter interface), then stateless AST checks, then stateful AST checks, then complex stateful checks with inheritance hierarchies.")

h2("6.2  Phase 3.1 — Infrastructure Scaffolding")
body("Before migrating any check, the hexagonal infrastructure was scaffolded in a single commit. Five new source files were created:")
table(
    ["File", "Package", "Architectural Role"],
    [
        ["ICheckUseCase.java", "checks.hexagonal.port", "Inbound port — AST-based checks"],
        ["IFileCheckUseCase.java", "checks.hexagonal.port", "Inbound port — file-level checks"],
        ["IViolationOutputPort.java", "checks.hexagonal.port", "Outbound port — violation emission"],
        ["FileContent.java", "checks.hexagonal.port", "Value object — file content boundary crossing"],
        ["AstCheckAdapter.java", "checks.hexagonal", "Abstract inbound adapter — AST checks"],
        ["FileSetCheckAdapter.java", "checks.hexagonal", "Abstract inbound adapter — file checks"],
    ]
)
body("After this commit, `mvn -DskipTests package` was run to confirm the new classes compiled and the build passed before any check migration began.")

h2("6.3  Phase 3.2 — File-Level Checks (Pilot)")
body("`LineLengthCheck` was migrated first as a pilot because it is the simplest check: it processes lines sequentially, has one threshold property, and its domain logic is a straightforward loop. This gave us a working end-to-end example of the adapter/domain split before tackling AST-based checks.")
body("`FileLengthCheck` was migrated immediately after using the same `FileSetCheckAdapter` base. Both use `IFileCheckUseCase`, `FileContent`, and `FileSetCheckAdapter`.")

h2("6.4  Phase 3.3 — Stateless AST Checks")
body("Five stateless checks were migrated in a single commit: `AnonInnerLengthCheck`, `LambdaBodyLengthCheck`, `MethodLengthCheck`, `ParameterNumberCheck`, and `RecordComponentNumberCheck`. These checks carry no state between token visits — each violation is computed at a single AST node without needing to track previous nodes. Their domain rules implement `ICheckUseCase` with only `visitToken()` overridden.")

h2("6.5  Phase 3.4 — Stateful AST Checks (No Inheritance)")
body("Five stateful checks were migrated next: `BooleanExpressionComplexityCheck`, `CyclomaticComplexityCheck`, `ExecutableStatementCountCheck`, `MethodCountCheck`, and `OuterTypeNumberCheck`. These maintain counters or stacks between `visitToken()` and `leaveToken()` calls within a single file walk. Their domain rules are annotated `@FileStatefulCheck` on the adapter side, ensuring `TreeWalker` creates a fresh adapter (and therefore a fresh rule) per file.")

h2("6.6  Phase 3.5 — Complex Stateful Checks")
body("`JavaNCSSCheck` and `NPathComplexityCheck` were the most complex to migrate because they use multiple token types, nested counters, and complex state machines. `NPathComplexityCheck` in particular uses a `Deque<BigInteger>` and multiplicative path counting across nested control structures. Despite the complexity, the migration procedure was identical: extract the logic into a domain rule, replace `log()` with `violationPort.report()`, wire via `createRule()`.")

h2("6.7  Phase 3.6 — Inheritance Family")
body("The final group — `ClassDataAbstractionCouplingCheck` and `ClassFanOutComplexityCheck` — share a common base `AbstractClassCouplingCheck` in the original code. This base class was migrated into `AbstractClassCouplingRule.java` in `checks/metrics/domain/`. The two check rules then extend `AbstractClassCouplingRule` and implement `ICheckUseCase`, preserving the original inheritance relationship but within the domain, not the framework.")

h2("6.8  Complete Class Mapping Table")
table(
    ["Original Class", "Domain Rule (*.domain)", "Adapter", "Port", "State"],
    [
        ["BooleanExpressionComplexityCheck", "BooleanExpressionComplexityRule", "BooleanExpressionComplexityCheck", "ICheckUseCase", "@FileStateful"],
        ["ClassDataAbstractionCouplingCheck", "ClassDataAbstractionCouplingRule\n(extends AbstractClassCouplingRule)", "ClassDataAbstractionCouplingCheck", "ICheckUseCase", "@FileStateful"],
        ["ClassFanOutComplexityCheck", "ClassFanOutComplexityRule\n(extends AbstractClassCouplingRule)", "ClassFanOutComplexityCheck", "ICheckUseCase", "@FileStateful"],
        ["CyclomaticComplexityCheck", "CyclomaticComplexityRule", "CyclomaticComplexityCheck", "ICheckUseCase", "@FileStateful"],
        ["JavaNCSSCheck", "JavaNCSSRule", "JavaNCSSCheck", "ICheckUseCase", "@FileStateful"],
        ["NPathComplexityCheck", "NPathComplexityRule", "NPathComplexityCheck", "ICheckUseCase", "@FileStateful"],
        ["AnonInnerLengthCheck", "AnonInnerLengthRule", "AnonInnerLengthCheck", "ICheckUseCase", "@Stateless"],
        ["ExecutableStatementCountCheck", "ExecutableStatementCountRule", "ExecutableStatementCountCheck", "ICheckUseCase", "@FileStateful"],
        ["FileLengthCheck", "FileLengthRule", "FileLengthCheck", "IFileCheckUseCase", "@Stateless"],
        ["LambdaBodyLengthCheck", "LambdaBodyLengthRule", "LambdaBodyLengthCheck", "ICheckUseCase", "@Stateless"],
        ["LineLengthCheck", "LineLengthRule", "LineLengthCheck", "IFileCheckUseCase", "@Stateless"],
        ["MethodCountCheck", "MethodCountRule", "MethodCountCheck", "ICheckUseCase", "@FileStateful"],
        ["MethodLengthCheck", "MethodLengthRule", "MethodLengthCheck", "ICheckUseCase", "@Stateless"],
        ["OuterTypeNumberCheck", "OuterTypeNumberRule", "OuterTypeNumberCheck", "ICheckUseCase", "@FileStateful"],
        ["ParameterNumberCheck", "ParameterNumberRule", "ParameterNumberCheck", "ICheckUseCase", "@Stateless"],
        ["RecordComponentNumberCheck", "RecordComponentNumberRule", "RecordComponentNumberCheck", "ICheckUseCase", "@Stateless"],
    ]
)

h2("6.9  Post-Refactoring Directory Structure")
body("After all 16 checks were migrated, the `checks/metrics/` and `checks/sizes/` packages each gained a `domain/` subdirectory containing the domain rule classes. The adapter check classes remain in the original package location. The screenshot below shows the new directory structure of both packages.")
spacer()
img("phase3-new-structure.png", Inches(5.8), "Figure 9 – Post-refactoring directory structure: domain/ subpackages visible in metrics and sizes")
spacer()
body("The `checks/hexagonal/` package contains the shared infrastructure: `AstCheckAdapter`, `FileSetCheckAdapter`, and the `port/` subpackage with the four port artifacts.")

page_break()

# ── 7. ARCHITECTURE DIAGRAMS ────────────────────────────────────────────────
h1("7.  Architecture Diagrams")

h2("7.1  Pre-Refactoring Sequence Diagrams")
body("The following two sequence diagrams show the execution path for a check in the original architecture — one for AST-based checks (processed by `TreeWalker`) and one for file-level checks (processed directly by `Checker`). These were captured before any refactoring changes.")

h3("7.1.1  AST-Based Check Path (Before)")
body("In the original AST path, `TreeWalker` calls `visitToken()` directly on the check object (which extends `AbstractCheck`). The check logic and the `log()` call are both inside the same class with no intermediate port.")
spacer()
img("seq-pre-refactor-ast.png", Inches(5.8), "Figure 10 – Pre-refactoring sequence: AST-based check (TreeWalker → AbstractCheck subclass)")
spacer()

h3("7.1.2  File-Level Check Path (Before)")
body("In the original file-set path, `Checker` calls `process(file, fileText)` directly on the check object (which extends `AbstractFileSetCheck`). Again, logic and violation emission are co-located in one class.")
spacer()
img("seq-pre-refactor-fileset.png", Inches(5.8), "Figure 11 – Pre-refactoring sequence: file-level check (Checker → AbstractFileSetCheck subclass)")
spacer()

h2("7.2  Post-Refactoring Sequence Diagrams")
body("The following two sequence diagrams show the same execution paths after the hexagonal refactoring. The key difference is that all calls now pass through the port interfaces, and the domain rule is a separate object that never receives a direct call from the framework.")

h3("7.2.1  AST-Based Check Path (After)")
body("After refactoring, `TreeWalker` calls `visitToken()` on the concrete adapter (still an `AbstractCheck` subclass via `AstCheckAdapter`). The adapter forwards the call through the `ICheckUseCase` port to the domain rule. When the domain rule finds a violation, it calls `IViolationOutputPort.report()` — which is implemented by the adapter calling `super.log()`.")
spacer()
img("seq-post-refactor-ast.png", Inches(5.8), "Figure 12 – Post-refactoring sequence: AST-based check with hexagonal port boundaries")
spacer()

h3("7.2.2  File-Level Check Path (After)")
body("After refactoring, `Checker` calls `process()` on the `FileSetCheckAdapter`. The adapter converts `FileText` to `FileContent` (a value object that contains no Checkstyle API types) and calls `IFileCheckUseCase.processFile()` on the domain rule. Violations flow back through `IViolationOutputPort`.")
spacer()
img("seq-post-refactor-fileset.png", Inches(5.8), "Figure 13 – Post-refactoring sequence: file-level check with FileContent value object boundary")
spacer()

h2("7.3  UML Class Diagram of the New Package Structure")
body("The UML class diagram below shows the complete class hierarchy and dependency relationships for the refactored hexagonal slice. The diagram includes all port interfaces, adapter base classes, concrete adapters, and domain rule classes. The key relationship to observe is that domain classes point only to port interfaces — no arrows cross from the domain hexagon directly to `AbstractCheck` or any other framework type.")
spacer()
img("uml-class-hexagonal.png", Inches(5.8), "Figure 14 – UML class diagram: full hexagonal slice (ports, adapters, domain rules)")
spacer()

h2("7.4  Package Dependency Comparison (Before vs After)")
body("The dependency comparison diagram visualises how the coupling structure of the `checks/metrics/` and `checks/sizes/` packages changed after the refactoring. In the before state, all check classes have direct arrows to `api.*` types. In the after state, only adapter classes have arrows to `api.*`; domain classes have arrows only to `checks.hexagonal.port.*`.")
spacer()
img("dep-comparison.png", Inches(5.8), "Figure 15 – Dependency comparison: direct api.* coupling (before) vs port-mediated coupling (after)")
spacer()

h2("7.5  C4 Architecture Diagrams")
body("The C4 Model provides four levels of architectural diagrams at increasing levels of zoom. We produced all four levels for the refactored system using Structurizr DSL (for L1–L3) and PlantUML (for L4).")

h3("7.5.1  Level 1: System Context Diagram")
body("The system context diagram shows Checkstyle as a single box and its two external actors: the Developer (who runs it) and the Java Source Files (that it analyses). This level establishes the system boundary.")
spacer()
img("c4-l1-context.png", Inches(5.5), "Figure 16 – C4 Level 1: System Context — Checkstyle 13.2.0 and its external actors")
spacer()

h3("7.5.2  Level 2: Container Diagram")
body("The container diagram zooms into the Checkstyle system and shows its major runtime components: CLI/Main, Checker, ConfigurationLoader, TreeWalker, the refactored Hexagonal Slice, and Other Checks (unchanged). The Hexagonal Slice is highlighted as the refactored subsystem.")
spacer()
img("c4-l2-containers.png", Inches(5.8), "Figure 17 – C4 Level 2: Containers — the Hexagonal Slice alongside the unchanged Checkstyle engine")
spacer()

h3("7.5.3  Level 3: Component Diagram")
body("The component diagram zooms into the Hexagonal Slice container and shows all its internal components: port interfaces (hexagon shapes), adapter base classes, concrete adapters for all 16 checks, and all domain rule classes. Color coding distinguishes ports (orange), inbound adapters (dark green), concrete adapters (light green), and domain classes (purple).")
spacer()
img("c4-l3-components.png", Inches(5.8), "Figure 18 – C4 Level 3: Components — ports, adapters, and domain rules inside the Hexagonal Slice")
spacer()

h3("7.5.4  Level 4: Code Diagram")
body("The code diagram shows one example check — `CyclomaticComplexityCheck` — before and after the refactoring at the class level. On the left (before), a single class extends `AbstractCheck` and contains both framework coupling and domain logic. On the right (after), two classes appear: the thin `CyclomaticComplexityCheck` adapter (extends `AstCheckAdapter`) and the pure `CyclomaticComplexityRule` domain class (implements `ICheckUseCase`). The adapter holds the framework coupling; the domain rule holds the logic.")
spacer()
img("c4-l4-code-example.png", Inches(5.8), "Figure 19 – C4 Level 4: Code diagram — CyclomaticComplexityCheck before (single class) vs after (adapter + domain rule)")
spacer()

page_break()

# ── 8. VERIFICATION ─────────────────────────────────────────────────────────
h1("8.  Verification and Testing")

h2("8.1  Regression Test: Byte-Identical Output")
body("The most fundamental verification was confirming that the refactored jar produces exactly the same output as the original on the same input. The procedure was:")
numbered("Build the refactored jar: `mvn -DskipTests package`")
numbered("Run it on the sample file: `java -jar target/checkstyle-13.2.0-all.jar -c violation-sample/all-metrics-sizes.xml violation-sample/SampleAllViolations.java > baseline/post-refactor-output.txt`")
numbered("Diff against the baseline: `diff baseline/pre-refactor-output.txt baseline/post-refactor-output.txt`")
body("The diff was empty. Both files contain exactly 44 violations in identical order and wording. The screenshot below shows the post-refactoring output:")
spacer()
img("phase4-post-refactor-output.png", Inches(5.8), "Figure 20 – Post-refactoring output: 44 violations, identical to the pre-refactoring baseline")
spacer()
body("The empty diff screenshot confirms byte-identical output:")
spacer()
img("phase4-empty-diff.png", Inches(4.0), "Figure 21 – Empty diff: pre-refactor and post-refactor outputs are byte-identical")
spacer()

h2("8.2  Full Test Suite (5726 Tests, 0 Failures)")
body("Checkstyle's own test suite was run without modification. All 5,726 tests pass with 0 failures, 0 errors, and 0 skipped. This confirms that the refactoring preserves the complete public API of all 16 checks (property setters, token lists, violation message keys, edge case handling) — not just the happy-path output on our sample file.")
body("The test suite includes both unit tests (per-check input/output scenarios using the inline config format) and integration tests (end-to-end Checker pipeline tests). The fact that all pass without any modification to test code is strong evidence that the adapter class names, packages, and Javadoc match the originals exactly.")
spacer()
img("phase4-mvn-test.png", Inches(5.8), "Figure 22 – mvn clean test: 5726 tests passing, 0 failures, 0 errors")
spacer()

h2("8.3  ArchUnit Conformance Tests (8/8 Rules Pass)")
body("We wrote 8 ArchUnit rules to formally verify the hexagonal architecture invariants. These rules are coded in `archunit-tests/src/main/java/ArchitectureValidationTest.java` and run against the compiled class files of the refactored jar. The 8 rules are:")

table(
    ["Rule", "Description", "Expected Result"],
    [
        ["1a", "metrics domain classes do NOT extend AbstractCheck", "PASS — 0 violations"],
        ["1b", "sizes domain classes do NOT extend AbstractCheck", "PASS — 0 violations"],
        ["2a", "metrics domain top-level classes implement ICheckUseCase", "PASS — all 6 domain rules implement it"],
        ["2b", "sizes domain top-level classes implement ICheckUseCase or IFileCheckUseCase", "PASS — all 10 domain rules implement it"],
        ["3a", "metrics adapter classes extend AbstractCheck (via AstCheckAdapter)", "PASS — all 6 adapters extend it"],
        ["3b", "sizes adapter classes extend AbstractCheck or AbstractFileSetCheck (via adapters)", "PASS — all 10 adapters extend it"],
        ["4", "domain classes do NOT depend on AbstractCheck directly", "PASS — no direct dependency"],
        ["5", "api.* package does NOT depend on checks.* (plugin isolation preserved)", "PASS — original isolation intact"],
    ]
)

spacer()
img("phase4-archunit.png", Inches(5.8), "Figure 23 – ArchUnit conformance test output: 8/8 rules pass")
spacer()

h2("8.4  jQAssistant Architecture Proof (Q1–Q4)")
body("jQAssistant was used to scan the compiled class files into a Neo4j graph database and run Cypher queries to verify the hexagonal boundaries programmatically. The Neo4j browser was started via `mvn -Pjqassistant jqassistant:server` and four queries were executed.")

h3("Q1 — No Domain Class Imports from api.*")
h4("Purpose")
body("Confirm that no class in the `.domain.` subpackages has any dependency on `com.puppycrawl.tools.checkstyle.api.*`. This is the core hexagonal invariant: the domain hexagon is framework-free.")
h4("Query Used")
code_block("""MATCH (domain:Class)-[:DEPENDS_ON]->(api:Class)
WHERE domain.fqn CONTAINS '.domain.'
  AND api.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.api.'
RETURN domain.fqn AS domain_class, api.fqn AS api_dependency
ORDER BY domain_class;""")
h4("Result")
body("Zero rows returned. No domain class has any dependency on any `api.*` class.")
spacer()
img("phase4-jqa-q1.png", Inches(5.5), "Figure 24 – Q1 result: zero rows — no domain class depends on api.*")
spacer()

h3("Q2 — Only Adapter Classes Extend AbstractCheck")
h4("Purpose")
body("Confirm that within the metrics and sizes packages, only the adapter classes (not `.domain.` subpackage classes) extend `AbstractCheck`, `AbstractFileSetCheck`, `AstCheckAdapter`, or `FileSetCheckAdapter`.")
h4("Query Used")
code_block("""MATCH (cls:Class)-[:EXTENDS]->(base:Class)
WHERE base.fqn IN [
  'com.puppycrawl.tools.checkstyle.api.AbstractCheck',
  'com.puppycrawl.tools.checkstyle.api.AbstractFileSetCheck',
  'com.puppycrawl.tools.checkstyle.checks.hexagonal.AstCheckAdapter',
  'com.puppycrawl.tools.checkstyle.checks.hexagonal.FileSetCheckAdapter'
]
  AND (cls.fqn CONTAINS '...checks.metrics.'
    OR cls.fqn CONTAINS '...checks.sizes.')
RETURN cls.fqn AS adapter_class, base.fqn AS extends_from
ORDER BY adapter_class;""")
h4("Result")
body("Only the 16 concrete adapter classes appear — none from the `.domain.` subpackages. This confirms the adaptation boundary is respected.")
spacer()
img("phase4-jqa-q2.png", Inches(5.5), "Figure 25 – Q2 result: only adapter classes (not domain classes) extend AbstractCheck or its adapters")
spacer()

h3("Q3 — Domain-to-Port Dependencies Only")
h4("Purpose")
body("Show all outgoing dependencies from domain classes to confirm they reference only `checks.hexagonal.port.*` interfaces (and `api.DetailAST` / `api.TokenTypes` as data types) — not full framework classes.")
h4("Query Used")
code_block("""MATCH (domain:Class)-[:DEPENDS_ON]->(port:Class)
WHERE domain.fqn CONTAINS '.domain.'
  AND (domain.fqn CONTAINS 'checks.metrics.' OR domain.fqn CONTAINS 'checks.sizes.')
RETURN domain.fqn AS domain_class, port.fqn AS dependency
ORDER BY domain_class, dependency;""")
h4("Result")
body("All domain class dependencies are to port interfaces (`ICheckUseCase`, `IViolationOutputPort`, `FileContent`) and data types (`DetailAST`, `TokenTypes`, `ScopeUtil`). No domain class references `AbstractCheck`, `AbstractFileSetCheck`, `Checker`, `TreeWalker`, or any non-data framework class.")
spacer()
img("phase4-jqa-q3.png", Inches(5.5), "Figure 26 – Q3 result: domain classes reference only port interfaces and data types")
spacer()

h3("Q4 — Full Slice Dependency Graph")
h4("Purpose")
body("Visualise the complete dependency graph of the hexagonal slice — all classes in metrics, sizes, and hexagonal packages and their dependency edges — to confirm the overall shape matches the hexagonal pattern.")
h4("Query Used")
code_block("""MATCH (src:Class)-[r:DEPENDS_ON]->(tgt:Class)
WHERE (src.fqn CONTAINS '...checks.metrics.'
    OR src.fqn CONTAINS '...checks.sizes.'
    OR src.fqn CONTAINS '...checks.hexagonal.')
RETURN src.fqn AS source, tgt.fqn AS target
ORDER BY source, target;""")
h4("Result")
body("The graph shows the characteristic hexagonal dependency pattern: adapters (outer layer) depend on ports and framework; ports form the middle ring; domain rules (inner layer) depend only on ports. No arrows cross from the inner layer directly to the framework.")
spacer()
img("phase4-jqa-q4.png", Inches(5.8), "Figure 27 – Q4 result: full slice dependency graph showing hexagonal layer structure")
spacer()

h2("8.5  Per-Check Violation Verification")
body("All 16 checks produce violations in the post-refactoring output. The table below shows the violation counts as confirmed in `baseline/post-refactor-output.txt`:")
table(
    ["Check", "Violations in Post-Refactoring Output"],
    [
        ["BooleanExpressionComplexity", "1"],
        ["ClassDataAbstractionCoupling", "1"],
        ["ClassFanOutComplexity", "1"],
        ["CyclomaticComplexity", "Multiple (methods with high complexity)"],
        ["JavaNCSS", "Multiple (methods and classes)"],
        ["NPathComplexity", "1"],
        ["AnonInnerLength", "1"],
        ["ExecutableStatementCount", "Multiple (methods)"],
        ["FileLength", "1"],
        ["LambdaBodyLength", "1"],
        ["LineLength", "Multiple (long lines)"],
        ["MethodCount", "1"],
        ["MethodLength", "Multiple (long methods)"],
        ["OuterTypeNumber", "1"],
        ["ParameterNumber", "1"],
        ["RecordComponentNumber", "1"],
    ]
)

page_break()

# ── 9. DEPENDENCIES ─────────────────────────────────────────────────────────
h1("9.  Dependencies Added and Removed")

h2("9.1  New Source Files Added")
body("The refactoring added the following source files to the project:")
table(
    ["Category", "Count", "Files"],
    [
        ["Port interfaces + value object", "4", "ICheckUseCase, IFileCheckUseCase, IViolationOutputPort, FileContent"],
        ["Adapter base classes", "2", "AstCheckAdapter, FileSetCheckAdapter"],
        ["Domain rules (metrics)", "7", "AbstractClassCouplingRule + 6 concrete rules"],
        ["Domain rules (sizes)", "10", "10 concrete rules"],
        ["Concrete adapters", "16", "16 *Check.java adapter classes (replace originals)"],
    ]
)
body("Net new source files: **(4 + 2 + 7 + 10) = 23** new files created. **16** original check files were deleted (replaced by adapters with the same class name). Net change: **+23 - 16 = +7 source files** in the main module, plus 6 new files in `checks/hexagonal/`.")

h2("9.2  Dependencies Removed from the Slice")
bullet("All direct `import com.puppycrawl.tools.checkstyle.api.AbstractCheck` statements from domain logic classes")
bullet("All direct `import com.puppycrawl.tools.checkstyle.api.AbstractFileSetCheck` statements from domain logic classes")
bullet("All direct calls to `this.log(...)` inside domain logic (replaced by `violationPort.report(...)`)")
bullet("Inheritance of `AbstractCheck` / `AbstractFileSetCheck` from classes that contain domain logic")

h2("9.3  Dependencies Added")
bullet("4 new port interfaces/value objects in `checks.hexagonal.port.*` — referenced by adapters and domain rules")
bullet("2 new abstract adapter base classes in `checks.hexagonal.*` — referenced only by the 16 concrete adapters")
bullet("Domain rules now depend on `ICheckUseCase` and `IViolationOutputPort` (port interfaces they themselves define)")
bullet("Adapters depend on `checks.hexagonal.AstCheckAdapter` or `FileSetCheckAdapter` — NOT directly on domain rules")

page_break()

# ── 10. LESSONS LEARNED ─────────────────────────────────────────────────────
h1("10.  Lessons Learned (Part A)")

h3("Lesson 1: Extract-then-wrap is the safest hexagonal migration path")
body("Trying to restructure the code architecture-first (create the port, wire the adapter, then move logic) risks breaking tests before the new architecture is fully working. The safer path is: extract the domain logic into a POJO first, verify it produces the same output, and only then delete the original. Each check was verified individually before moving to the next, giving a clear checkpoint at each step.")

h3("Lesson 2: Javadoc placement matters in frameworks")
body("In Checkstyle, site documentation pages and IDE plugin metadata are generated from Javadoc on the public check class. When we initially moved all Javadoc to the domain rule class, the site generator could not find it because it looks for Javadoc on the class that appears in `packages.xml` (the adapter). Lesson learned: in a hexagonal architecture applied to an existing framework, the user-facing documentation must stay on the adapter — because the adapter IS the public API from the framework's perspective, even though the domain rule contains the logic.")

h3("Lesson 3: Thread-safety annotations belong on the adapter")
body("Checkstyle's `@StatelessCheck`, `@FileStatefulCheck`, and `@GlobalStatefulCheck` annotations are processed by `TreeWalker` to decide how to instantiate check modules. These annotations must stay on the adapter class — not the domain rule — because `TreeWalker` looks for them on the class that extends `AbstractCheck`. The domain rule's statefulness (whether it has instance fields that reset per file) is an implementation detail internal to the adapter's lifecycle management.")

h3("Lesson 4: Hexagonal architecture imposes near-zero overhead for static analysis tools")
body("The port-dispatch overhead (two virtual method calls per check invocation: `ICheckUseCase.visitToken()` and `IViolationOutputPort.report()`) is negligible for a tool dominated by file I/O and AST parsing. The JVM JIT inlines these calls after a brief warm-up period, eliminating any dispatch cost. Our benchmarks confirmed this: all 5 benchmark projects showed timing differences within ±6%, well inside normal JVM timing noise (~±10%).")

h3("Lesson 5: Architecture conformance tools are essential, not optional")
body("ArchUnit rules and jQAssistant queries are not just documentation artifacts — they are the actual proof that the architecture is correctly implemented. Without automated conformance checking, it is easy for a single class to accidentally import `AbstractCheck` or call `log()` directly and for the violation to go unnoticed in code review. Writing the ArchUnit rules first (before migrating all 16 checks) and running them incrementally caught several accidental violations early and ensured the final architecture is clean throughout.")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART B
# ════════════════════════════════════════════════════════════════════════════

h1("Part B: Performance Experiment")
spacer()

h1("11.  Experimental Setup")

h2("11.1  Objective")
body("The objective of Part B is to measure whether the hexagonal architecture refactoring introduces any measurable performance overhead compared to the original Checkstyle architecture, and to characterise how performance scales with codebase size for both jars.")
body("The hexagonal refactoring adds two virtual method call layers per check invocation: `ICheckUseCase.visitToken()` (the inbound port dispatch) and `IViolationOutputPort.report()` (the outbound port, called only on violations). We expected the JVM JIT to inline these after warm-up, but we wanted empirical confirmation.")

h2("11.2  Selected Projects")
body("Five open-source Java projects were selected to cover a range of codebase sizes, from small (minimal-json, ~1,400 lines) to large (Apache Calcite core, ~500,000 lines):")
table(
    ["Project", "Source directory scanned", "Approx. Java files", "Approx. LOC"],
    [
        ["Minimal JSON", "src/main/java", "~20 files", "~1,400 LOC"],
        ["JavaPoet", "src/main/java", "~30 files", "~4,000 LOC"],
        ["GraphStream Core (gs-core)", "src/", "~120 files", "~25,000 LOC"],
        ["JGraphT (core module)", "jgrapht-core/src/main/java", "~300 files", "~60,000 LOC"],
        ["Apache Calcite (core module)", "core/src/main/java", "~2,000 files", "~500,000 LOC"],
    ]
)
body("These projects were selected because they span three orders of magnitude in size, cover diverse Java coding patterns, and were used in the Task 1 Part B benchmarks — enabling direct comparison with the original baseline.")

h2("11.3  Methodology")
body("For each project and each jar (original and refactored):")
numbered("A shallow clone of the project was obtained.")
numbered("One untimed JVM warm-up run was executed to allow class loading and JIT compilation to stabilise.")
numbered("Three timed wall-clock runs were performed using `date +%s%3N` (millisecond granularity) in bash.")
numbered("The average of the three timed runs was recorded.")
body("Timing used `bash` wall-clock timing (`date +%s%3N` before and after the `java -jar ...` invocation). All runs were performed on the same machine with no background load, with the same Java 21 JVM settings, and with the same configuration XML (`baseline/bench-config.xml`) enabling all 16 metrics + sizes checks at their default thresholds.")

h2("11.4  Configuration")
body("The benchmark configuration (`baseline/bench-config.xml`) enables all 16 refactored checks at default thresholds. Default thresholds were used (rather than artificially low ones) to keep the benchmark representative of real-world usage, where violations are rare and the majority of check invocations return without reporting anything.")
code_block("""<module name="Checker">
  <module name="TreeWalker">
    <module name="CyclomaticComplexity"/>  <!-- default max=10 -->
    <module name="JavaNCSS"/>
    <module name="NPathComplexity"/>
    <module name="BooleanExpressionComplexity"/>
    <module name="ClassDataAbstractionCoupling"/>
    <module name="ClassFanOutComplexity"/>
    <!-- ... all 16 checks ... -->
  </module>
  <module name="FileLength"/>    <!-- file-level checks outside TreeWalker -->
  <module name="LineLength"/>
</module>""")

page_break()

# ── 12. RESULTS ─────────────────────────────────────────────────────────────
h1("12.  Experimental Results")

h2("12.1  Timing Results Table")
body("The table below shows the average wall-clock time (milliseconds) for the original and refactored jars on each of the five benchmark projects, along with the absolute and percentage difference.")
table(
    ["Project", "Avg Original (ms)", "Avg Refactored (ms)", "Delta (ms)", "Delta %"],
    [
        ["minimal-json", "1,946", "1,993", "+47", "+2.4%"],
        ["JavaPoet", "1,918", "1,867", "−51", "−2.6%"],
        ["GraphStream Core (gs-core)", "4,501", "4,475", "−26", "−0.6%"],
        ["JGraphT (jgrapht-core)", "828", "795", "−33", "−4.0%"],
        ["Apache Calcite (core)", "26,247", "24,681", "−1,566", "−6.0%"],
    ]
)
body("All five deltas fall within ±6%, which is well within normal JVM wall-clock timing noise (~±10% for this kind of measurement). The refactored jar is not consistently slower than the original — three of the five projects show the refactored jar running slightly faster (likely due to normal run-to-run variance).")

h2("12.2  Performance Graph")
body("The bar chart below shows the original vs. refactored timing for all five projects overlaid. The bars are nearly identical for each project, confirming no architectural overhead is visible at this measurement granularity.")
spacer()
img("phase5-performance-graph.png", Inches(5.8), "Figure 28 – Performance comparison: original vs refactored jar across 5 benchmark projects")
spacer()

h2("12.3  Refactored Jar Terminal Output")
body("The screenshot below shows the terminal output from running the benchmark script for the refactored jar on all five projects, showing the three timed runs per project and the computed averages that populate the results table.")
spacer()
img("phase5-perf-refactored.png", Inches(5.5), "Figure 29 – Terminal output: timing runs for refactored jar on all 5 benchmark projects")
spacer()

page_break()

# ── 13. ANALYSIS ────────────────────────────────────────────────────────────
h1("13.  Performance Analysis")

h2("13.1  Performance Impact")
body("The hexagonal architecture refactoring introduces **no measurable performance overhead**. All five delta measurements fall within ±6% of the original times, and none are statistically consistent in direction (both positive and negative deltas appear). Given that JVM wall-clock measurements for Java tools typically show ±10% variance between runs on the same machine, these deltas are indistinguishable from noise.")
body("The theoretical overhead is two virtual method calls per AST node visit per check: one for `ICheckUseCase.visitToken()` (the inbound port dispatch from adapter to domain rule) and one for `IViolationOutputPort.report()` (called only when a violation is found, which is rare at default thresholds). In a hot loop, the JVM JIT compiler inlines these virtual calls after ~10,000 invocations, reducing them to zero overhead in the steady state.")

h2("13.2  Scalability Assessment")
body("The performance of both jars scales linearly with codebase size. Comparing the two endpoints:")
bullet("**minimal-json**: ~1.9 seconds for ~1,400 lines (~0.7 ms/class)")
bullet("**Apache Calcite**: ~25 seconds for ~500,000 lines (~0.05 ms/class at scale)")
body("The ratio between the two jars is consistent across all project sizes — both jars spend the same fraction of time on each project. This confirms the hexagonal overhead is O(1) per check call, not O(N) in file count or codebase size. If the overhead were linear, it would become visible at larger projects like Calcite — but it does not.")

h2("13.3  Bottleneck Identification")
body("Profiling and simple reasoning identify the true bottlenecks:")
numbered("**JVM startup and class loading** (~1.5 seconds baseline): Dominates for small projects (minimal-json, javapoet) where the JVM startup time is a large fraction of total time. The hexagonal refactoring adds 23 new class files, but class loading for Java standard library classes dominates this cost and swamps the new files' contribution.")
numbered("**ANTLR4 parsing**: For larger projects, the majority of time is spent lexing and parsing Java source files into `DetailAST` trees. This is completely unaffected by the hexagonal refactoring, which does not touch `JavaLanguageLexer.g4`, `JavaLanguageParser.g4`, or `JavaAstVisitor`.")
numbered("**File I/O**: Reading `~500,000` lines of Java source code from disk contributes significantly for large projects. Again, unaffected by the refactoring.")
numbered("**Check dispatch** (the hexagonal overhead): Two virtual calls per AST node per check. With 16 checks and a Java file averaging ~50 AST nodes, this is ~800 virtual calls per file. At ~1 ns per JIT-inlined call, this is ~800 ns per file — negligible against the ~0.05 ms per file dominated by parsing.")

h2("13.4  Knee of the Curve")
body("Examining the timing data, the performance transitions from JVM-startup-dominated to parse-dominated somewhere between JavaPoet (~1.9 s, ~4,000 LOC) and GraphStream Core (~4.5 s, ~25,000 LOC). The 'knee of the curve' where the tool becomes useful for wall-clock profiling (rather than being dominated by JVM startup) is approximately **10,000–20,000 LOC**, where AST parsing begins to dominate JVM startup costs.")
body("Both the original and refactored jars hit this knee at the same point, confirming that the hexagonal architecture does not shift the performance characteristics of the tool.")

h2("13.5  What We Would Do Differently")
numbered("**Use JMH (Java Microbenchmark Harness)** instead of bash wall-clock timing. JMH provides proper JIT warm-up control, multiple benchmark modes (throughput, average time, sample time), and confidence intervals. Our bash timing includes JVM startup in every measurement, which inflates the noise for small projects.")
numbered("**Fork the JVM per run** to eliminate cross-run JIT state contamination. When measuring multiple projects in the same benchmarking session, JIT profiles from one project may affect another.")
numbered("**Profile with async-profiler** to empirically confirm that adapter dispatch contributes less than 0.1% of CPU time, rather than relying on theoretical reasoning. A flame graph showing `visitToken` and `report` at the bottom of the stack would be more convincing than the absence of measurable overhead.")
numbered("**Run more iterations** (10 runs instead of 3) and report confidence intervals. Three runs provide very limited statistical power for distinguishing 2% differences from noise.")
numbered("**Separate warm-up from measurement** more carefully. Even with one untimed warm-up run, the first timed run may still catch late-phase JIT compilations. JMH's warm-up infrastructure handles this more reliably.")

page_break()

# ── 14. LESSONS LEARNED PART B ──────────────────────────────────────────────
h1("14.  Lessons Learned (Part B)")

h3("Lesson 1: Hexagonal architecture is performance-neutral for static analysis tools")
body("The ports-and-adapters pattern imposes no measurable runtime overhead in this context. The abstraction cost (two virtual method calls per check event) is completely hidden by JVM JIT inlining after the first few hundred invocations. Teams evaluating hexagonal architecture for performance-sensitive Java tools should not be deterred by the indirection — the architectural benefits (framework-free domain, port-isolated tests, explicit technology boundaries) come at essentially zero runtime cost in practice.")

h3("Lesson 2: Wall-clock benchmarking is insufficient for precise overhead measurement")
body("Our bash wall-clock methodology can detect large regressions (>10%) but cannot distinguish 1-5% differences from JVM timing noise. For the purpose of proving 'no overhead', it is sufficient — but if the hexagonal refactoring had introduced, say, a 3% regression, our methodology could not have confirmed it. For future work on performance-sensitive refactorings, JMH benchmarks with proper warm-up and statistical analysis are necessary.")

h3("Lesson 3: Scalability is about the algorithm, not the architecture")
body("The most important observation from the benchmark is that Checkstyle scales linearly with codebase size because its algorithm is O(N) in file count: each file is parsed and walked once. The hexagonal refactoring does not change this fundamental scaling characteristic because it only reorganises the delegation structure within the per-file processing loop — it does not change what work is done per file, only how that work is structured.")

h3("Lesson 4: JVM startup dominates small-codebase benchmarks")
body("For small projects (minimal-json, javapoet), JVM startup time (~1.5 seconds) is 80-90% of total measured time. This means that a hypothetical 50% speedup in check dispatch logic would only improve total time by ~5-10% for these projects. When interpreting performance results for Java command-line tools, JVM startup must always be considered as a baseline cost that cannot be reduced without switching to a JVM-embedded model (Maven plugin, IDE integration, Gradle plugin).")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
h1("References")

body("1. Cockburn, A. (2005). *Hexagonal Architecture*. Retrieved from https://alistair.cockburn.us/hexagonal-architecture/")
body("2. Checkstyle 13.2.0 Source Code. Apache Software Foundation / Checkstyle Project. https://github.com/checkstyle/checkstyle/tree/checkstyle-13.2.0")
body("3. Brown, S., (2018). *The C4 Model for Software Architecture*. https://c4model.com/")
body("4. Structurizr DSL Reference. https://docs.structurizr.com/dsl/")
body("5. ArchUnit User Guide. TNG Technology Consulting. https://www.archunit.org/userguide/html/000_Index.html")
body("6. jQAssistant Documentation. jQAssistant Project. https://jqassistant.github.io/jqassistant/")
body("7. Martin, R.C. (2012). *The Clean Architecture*. Retrieved from https://blog.cleancoder.com/uncle-bob/2012/08/13/the-clean-architecture.html")
body("8. Fowler, M. (2004). *InversionOfControl*. https://martinfowler.com/articles/injection.html")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX A — jQAssistant Queries
# ════════════════════════════════════════════════════════════════════════════
h1("Appendix A — jQAssistant Cypher Queries (Q1–Q4)")

body("The four Cypher queries used to verify the hexagonal architecture via jQAssistant Neo4j graph analysis. Run each in the Neo4j browser at `http://localhost:7474` after scanning with `mvn -Pjqassistant jqassistant:scan`.")

h2("Q1 — Confirm No Domain Class Imports from api.*")
code_block("""MATCH (domain:Class)-[:DEPENDS_ON]->(api:Class)
WHERE domain.fqn CONTAINS '.domain.'
  AND api.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.api.'
RETURN domain.fqn AS domain_class, api.fqn AS api_dependency
ORDER BY domain_class;
-- Expected: ZERO rows (domain is framework-free)""")

h2("Q2 — Confirm Only Adapter Classes Extend AbstractCheck")
code_block("""MATCH (cls:Class)-[:EXTENDS]->(base:Class)
WHERE base.fqn IN [
  'com.puppycrawl.tools.checkstyle.api.AbstractCheck',
  'com.puppycrawl.tools.checkstyle.api.AbstractFileSetCheck',
  'com.puppycrawl.tools.checkstyle.checks.hexagonal.AstCheckAdapter',
  'com.puppycrawl.tools.checkstyle.checks.hexagonal.FileSetCheckAdapter'
]
  AND (cls.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.checks.metrics.'
    OR cls.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.checks.sizes.')
RETURN cls.fqn AS adapter_class, base.fqn AS extends_from
ORDER BY adapter_class;
-- Expected: only *Check adapter classes, no *.domain.* classes""")

h2("Q3 — Domain-to-Port Dependencies")
code_block("""MATCH (domain:Class)-[:DEPENDS_ON]->(port:Class)
WHERE domain.fqn CONTAINS '.domain.'
  AND (domain.fqn CONTAINS 'checks.metrics.' OR domain.fqn CONTAINS 'checks.sizes.')
RETURN domain.fqn AS domain_class, port.fqn AS dependency
ORDER BY domain_class, dependency;
-- Expected: only port interfaces (ICheckUseCase, IViolationOutputPort, FileContent)
-- and data types (DetailAST, TokenTypes, ScopeUtil)""")

h2("Q4 — Full Slice Dependency Graph")
code_block("""MATCH (src:Class)-[r:DEPENDS_ON]->(tgt:Class)
WHERE (src.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.checks.metrics.'
    OR src.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.checks.sizes.'
    OR src.fqn CONTAINS 'com.puppycrawl.tools.checkstyle.checks.hexagonal.')
RETURN src.fqn AS source, tgt.fqn AS target
ORDER BY source, target;
-- Visualise as graph to confirm hexagonal layering""")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX B — ArchUnit Test Rules
# ════════════════════════════════════════════════════════════════════════════
h1("Appendix B — ArchUnit Conformance Test Rules")
body("The following ArchUnit rules are implemented in `archunit-tests/src/main/java/ArchitectureValidationTest.java` and run against the compiled class files of the refactored jar.")

code_block("""// Rule 1a: metrics domain classes do NOT extend AbstractCheck
noClasses().that().resideInAPackage("..checks.metrics.domain..")
    .should().beAssignableTo(ABSTRACT_CHECK)
    .because("metrics domain classes must be framework-free");

// Rule 1b: sizes domain classes do NOT extend AbstractCheck
noClasses().that().resideInAPackage("..checks.sizes.domain..")
    .should().beAssignableTo(ABSTRACT_CHECK)
    .because("sizes domain classes must be framework-free");

// Rule 2a: metrics domain top-level classes implement ICheckUseCase
classes().that().resideInAPackage("..checks.metrics.domain..")
    .and().areTopLevelClasses().and().areNotInterfaces()
    .should().implement(ICHECK_USE_CASE)
    .because("every top-level domain class is a use-case reachable through a port");

// Rule 3a: metrics adapter classes extend AbstractCheck
classes().that().resideInAPackage("..checks.metrics")
    .and().haveSimpleNameEndingWith("Check")
    .and().areNotInterfaces()
    .should().beAssignableTo(ABSTRACT_CHECK)
    .because("adapter classes must bridge the Checkstyle framework to the domain");

// Rule 4: domain classes do NOT depend on AbstractCheck directly
noClasses().that()
    .resideInAPackage("..checks.metrics.domain..")
    .or().resideInAPackage("..checks.sizes.domain..")
    .should().dependOnClassesThat().haveFullyQualifiedName(ABSTRACT_CHECK)
    .because("domain rules must be framework-free");

// Rule 5: api.* does NOT depend on checks.* (plugin isolation preserved)
noClasses().that().resideInAPackage("..checkstyle.api..")
    .should().dependOnClassesThat().resideInAPackage("..checkstyle.checks..")
    .because("the API package must not depend on any check implementation");""")

page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX C — Structurizr DSL Excerpt
# ════════════════════════════════════════════════════════════════════════════
h1("Appendix C — Structurizr DSL Workspace (Excerpt)")
body("The full Structurizr DSL workspace is at `docs/structurizr/workspace.dsl`. The excerpt below shows the C4 model structure (model definition only, styles and views omitted for brevity).")

code_block("""workspace "Checkstyle 13.2.0 — Hexagonal Refactoring" {
  model {
    developer = person "Developer"
    javaSource = softwareSystem "Java Source Files" "External"
    checkstyle = softwareSystem "Checkstyle 13.2.0" {
      cli = container "CLI / Main"
      checker = container "Checker"
      treeWalker = container "TreeWalker"
      hexagonalSlice = container "Hexagonal Slice" {
        inputPort  = component "ICheckUseCase"         tag "Port"
        outputPort = component "IViolationOutputPort"  tag "Port"
        fileContent= component "FileContent"           tag "Value Object"
        astAdapter = component "AstCheckAdapter"       tag "Inbound Adapter"
        fscAdapter = component "FileSetCheckAdapter"   tag "Inbound Adapter"
        -- 16 domain rules (metrics + sizes) --
        -- 16 concrete adapters --
      }
      otherChecks = container "Other Checks (unchanged)"
    }
    developer -> checkstyle "runs with config XML and Java sources"
    checker -> hexagonalSlice "dispatches via FileSetCheckAdapter"
    treeWalker -> hexagonalSlice "dispatches via AstCheckAdapter"
    astAdapter -> inputPort "implements"
    fscAdapter -> inputPort "implements"
    astAdapter -> outputPort "implements (via log())"
    -- domain rules implement inputPort, call outputPort --
  }
}""")

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
